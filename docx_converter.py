#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Word文档转换器模块
支持DOCX和DOC文件转换为Markdown
"""

import os
import sys
from pathlib import Path
from typing import Optional, Tuple
import subprocess
import shutil


def convert_docx_to_markdown(docx_path: Path, out_dir: Path) -> Tuple[bool, str, Optional[Path]]:
    """
    将Word文档转换为Markdown
    
    参数:
        docx_path: Word文档路径
        out_dir: 输出目录
    
    返回:
        (success, message, output_file)
    """
    try:
        # 确保参数是Path对象
        docx_path_obj = Path(docx_path) if isinstance(docx_path, str) else docx_path
        out_dir_obj = Path(out_dir) if isinstance(out_dir, str) else out_dir
        
        # 创建输出目录
        out_dir_obj.mkdir(parents=True, exist_ok=True)
        
        # 方法1: 尝试使用pandoc（首选）
        if shutil.which("pandoc") is not None:
            output_file = out_dir_obj / f"{docx_path_obj.stem}.md"
            cmd = [
                "pandoc",
                "-s", str(docx_path_obj),
                "-t", "markdown",
                "-o", str(output_file)
            ]
            
            try:
                result = subprocess.run(
                    cmd,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    text=True,
                    timeout=30
                )
                
                if result.returncode == 0 and output_file.exists():
                    return True, f"使用pandoc转换成功", output_file
                else:
                    error_msg = result.stderr.strip() if result.stderr else "未知错误"
                    return False, f"pandoc转换失败: {error_msg}", None
            except subprocess.TimeoutExpired:
                return False, "pandoc转换超时", None
            except Exception as e:
                return False, f"pandoc执行异常: {e}", None
        
        # 方法2: 尝试使用python-docx
        try:
            import docx
            
            doc = docx.Document(str(docx_path_obj))
            output_file = out_dir_obj / f"{docx_path_obj.stem}.md"
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(f"# {docx_path_obj.stem}\n\n")
                
                for para in doc.paragraphs:
                    text = para.text.strip()
                    if text:
                        # 简单的格式检测
                        if para.style.name.startswith('Heading'):
                            level = 1
                            if para.style.name == 'Heading 2':
                                level = 2
                            elif para.style.name == 'Heading 3':
                                level = 3
                            f.write(f"{'#' * level} {text}\n\n")
                        else:
                            f.write(f"{text}\n\n")
                
                # 处理表格
                for table in doc.tables:
                    f.write("\n|")
                    # 简单的表格处理
                    for row in table.rows:
                        row_text = "|".join([cell.text.strip() for cell in row.cells])
                        f.write(f"{row_text}|\n")
                    f.write("\n")
            
            return True, "使用python-docx转换成功", output_file
            
        except ImportError:
            # python-docx不可用
            pass
        except Exception as e:
            # python-docx处理出错
            pass
        
        # 方法3: 尝试使用antiword（针对.doc文件）
        if docx_path_obj.suffix.lower() == '.doc' and shutil.which("antiword") is not None:
            output_file = out_dir_obj / f"{docx_path_obj.stem}.txt"
            cmd = ["antiword", str(docx_path_obj)]
            
            try:
                with open(output_file, 'w', encoding='utf-8') as f:
                    result = subprocess.run(
                        cmd,
                        stdout=f,
                        stderr=subprocess.PIPE,
                        text=True,
                        timeout=30
                    )
                
                if result.returncode == 0 and output_file.exists():
                    # 将.txt转换为.md
                    md_file = out_dir_obj / f"{docx_path_obj.stem}.md"
                    with open(output_file, 'r', encoding='utf-8') as f_in, \
                         open(md_file, 'w', encoding='utf-8') as f_out:
                        f_out.write(f"# {docx_path_obj.stem}\n\n")
                        f_out.write("```text\n")
                        f_out.write(f_in.read())
                        f_out.write("\n```\n")
                    
                    # 删除临时.txt文件
                    output_file.unlink(missing_ok=True)
                    
                    return True, "使用antiword转换成功", md_file
                else:
                    error_msg = result.stderr.strip() if result.stderr else "未知错误"
                    return False, f"antiword转换失败: {error_msg}", None
            except subprocess.TimeoutExpired:
                return False, "antiword转换超时", None
            except Exception as e:
                return False, f"antiword执行异常: {e}", None
        
        # 方法4: 最后尝试使用文本提取
        try:
            # 尝试使用catdoc或其他文本提取工具
            output_file = out_dir_obj / f"{docx_path_obj.stem}.md"
            
            # 简单的文本提取（作为最后手段）
            import zipfile
            if docx_path_obj.suffix.lower() == '.docx':
                # DOCX文件实际上是ZIP文件
                with zipfile.ZipFile(docx_path_obj, 'r') as zip_ref:
                    # 查找文档内容
                    if 'word/document.xml' in zip_ref.namelist():
                        import xml.etree.ElementTree as ET
                        from xml.etree.ElementTree import QName
                        
                        with zip_ref.open('word/document.xml') as xml_file:
                            tree = ET.parse(xml_file)
                            root = tree.getroot()
                            
                            # 简单的文本提取
                            namespaces = {
                                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                            }
                            
                            texts = []
                            for elem in root.iter():
                                if elem.tag.endswith('}t'):  # 文本元素
                                    if elem.text:
                                        texts.append(elem.text)
                            
                            with open(output_file, 'w', encoding='utf-8') as f:
                                f.write(f"# {docx_path_obj.stem}\n\n")
                                f.write("\n".join(texts))
                            
                            return True, "使用简单文本提取转换成功", output_file
            
            return False, "没有可用的Word文档转换工具", None
            
        except Exception as e:
            return False, f"文本提取失败: {e}", None
            
    except Exception as e:
        return False, f"转换过程异常: {e}", None


def ensure_word_converter_exists() -> bool:
    """
    检查是否有可用的Word文档转换工具
    
    返回:
        bool: 是否有可用的转换工具
    """
    # 检查pandoc
    if shutil.which("pandoc") is not None:
        return True
    
    # 检查python-docx
    try:
        import docx
        return True
    except ImportError:
        pass
    
    # 检查antiword
    if shutil.which("antiword") is not None:
        return True
    
    # 检查catdoc
    if shutil.which("catdoc") is not None:
        return True
    
    return False


def get_word_converter_help() -> str:
    """
    获取Word文档转换工具的安装帮助
    """
    help_text = """
没有找到可用的Word文档转换工具。请安装以下之一：

1. pandoc（推荐）:
   - macOS: brew install pandoc
   - Ubuntu/Debian: sudo apt-get install pandoc
   - Windows: 从 https://pandoc.org/installing.html 下载

2. python-docx（Python库）:
   pip install python-docx

3. antiword（针对.doc文件）:
   - macOS: brew install antiword
   - Ubuntu/Debian: sudo apt-get install antiword

4. catdoc（文本提取）:
   - macOS: brew install catdoc
   - Ubuntu/Debian: sudo apt-get install catdoc
"""
    return help_text


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("用法: python docx_converter.py <Word文档> <输出目录>")
        sys.exit(1)
    
    docx_path = Path(sys.argv[1])
    out_dir = Path(sys.argv[2])
    
    if not docx_path.exists():
        print(f"错误: Word文档不存在: {docx_path}")
        sys.exit(1)
    
    success, message, output_file = convert_docx_to_markdown(docx_path, out_dir)
    
    if success:
        print(f"转换成功: {message}")
        print(f"输出文件: {output_file}")
    else:
        print(f"转换失败: {message}")
        if not ensure_word_converter_exists():
            print(get_word_converter_help())
        sys.exit(1)